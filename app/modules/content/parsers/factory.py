import io
import re

from app.modules.content.parsers.base import BaseParser, ParsedSection

# ---------------------------------------------------------------------------
# PDF parser (pypdf)
# ---------------------------------------------------------------------------


class PDFParser(BaseParser):
    """Extract sections from a PDF by splitting on heading-like lines."""

    _HEADING_RE = re.compile(r"^(?P<num>(?:\d+\.)+\d*)\s+(?P<title>.+)$")

    def parse(self, file_bytes: bytes) -> list[ParsedSection]:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        sections: list[ParsedSection] = []
        current: ParsedSection | None = None
        buffer: list[str] = []
        ordinal = 0

        def _flush(sec: ParsedSection, lines: list[str]) -> None:
            sec.content_markdown = "\n".join(lines).strip()

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            for raw_line in text.splitlines():
                line = raw_line.strip()
                if not line:
                    continue
                m = self._HEADING_RE.match(line)
                if m:
                    if current is not None:
                        _flush(current, buffer)
                    ordinal += 1
                    current = ParsedSection(
                        section_number=m.group("num").rstrip("."),
                        title=m.group("title").strip(),
                        content_markdown="",
                        page_number=page_num,
                        ordinal=ordinal,
                    )
                    sections.append(current)
                    buffer = []
                else:
                    if current is None:
                        # Text before first heading → preamble section
                        ordinal += 1
                        current = ParsedSection(
                            section_number="0",
                            title="Preamble",
                            content_markdown="",
                            page_number=page_num,
                            ordinal=ordinal,
                        )
                        sections.append(current)
                    buffer.append(line)

        if current is not None:
            _flush(current, buffer)

        if not sections:
            # Fallback: single section with all text
            full_text = "\n".join((page.extract_text() or "") for page in reader.pages)
            sections.append(
                ParsedSection(
                    section_number="1",
                    title="Document",
                    content_markdown=full_text.strip(),
                    page_number=1,
                    ordinal=1,
                )
            )

        return sections


# ---------------------------------------------------------------------------
# DOCX parser (python-docx)
# ---------------------------------------------------------------------------


class DOCXParser(BaseParser):
    """Extract sections from a DOCX using Heading styles as section boundaries."""

    def parse(self, file_bytes: bytes) -> list[ParsedSection]:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        sections: list[ParsedSection] = []
        current: ParsedSection | None = None
        buffer: list[str] = []
        ordinal = 0

        def _flush(sec: ParsedSection, lines: list[str]) -> None:
            sec.content_markdown = "\n\n".join(lines).strip()

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                if current is not None:
                    _flush(current, buffer)
                ordinal += 1
                current = ParsedSection(
                    section_number=str(ordinal),
                    title=text,
                    content_markdown="",
                    page_number=None,
                    ordinal=ordinal,
                )
                sections.append(current)
                buffer = []
            else:
                if current is None:
                    ordinal += 1
                    current = ParsedSection(
                        section_number="0",
                        title="Preamble",
                        content_markdown="",
                        page_number=None,
                        ordinal=ordinal,
                    )
                    sections.append(current)
                buffer.append(text)

        if current is not None:
            _flush(current, buffer)

        if not sections:
            full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            sections.append(
                ParsedSection(
                    section_number="1",
                    title="Document",
                    content_markdown=full_text.strip(),
                    page_number=None,
                    ordinal=1,
                )
            )

        return sections


# ---------------------------------------------------------------------------
# Factory registry
# ---------------------------------------------------------------------------


class ParserFactory:
    _parsers: dict[str, type[BaseParser]] = {}

    @classmethod
    def register(cls, source_type: str, parser_cls: type[BaseParser]) -> None:
        cls._parsers[source_type] = parser_cls

    @classmethod
    def get_parser(cls, source_type: str) -> BaseParser:
        parser_cls = cls._parsers.get(source_type)
        if not parser_cls:
            raise ValueError(f"No parser registered for source type: {source_type}")
        return parser_cls()


# All aviation document types use the PDF parser by default;
# DOCX is the fallback for syllabus/SOP that may arrive as Word docs.
ParserFactory.register("fcom", PDFParser)
ParserFactory.register("qrh", PDFParser)
ParserFactory.register("amm", PDFParser)
ParserFactory.register("sop", DOCXParser)
ParserFactory.register("syllabus", DOCXParser)
